#!/usr/bin/env python3
"""
parse_docx.py
-------------
Parses DOCX files to extract text content from paragraphs and tables.
Each paragraph and table row is saved as a separate chunk.
The output is a list of dictionaries with:
  - chunk_id: Unique identifier.
  - modality: "text".
  - content: Extracted text.
  - metadata: Additional details like file name, paragraph index, table index, etc.
"""

import os
from docx import Document

def process_docx_file(file_path: str):
    """
    Extracts text from a DOCX file and returns a list of text chunks.

    :param file_path: Path to the DOCX file.
    :return: List of dictionaries representing text chunks.
    """
    chunks = []

    try:
        # Load the DOCX document.
        doc = Document(file_path)

        # Process each paragraph in the document.
        for p_index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                chunk_id = f"{os.path.basename(file_path)}_par_{p_index}"
                chunks.append({
                    "chunk_id": chunk_id,
                    "modality": "text",
                    "content": text,
                    "metadata": {
                        "file_name": os.path.basename(file_path),
                        "paragraph_index": p_index
                    }
                })

        # Process each table in the document.
        for t_index, table in enumerate(doc.tables):
            for r_index, row in enumerate(table.rows):
                # Concatenate cell text from each row, separated by commas.
                row_text = ", ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    chunk_id = f"{os.path.basename(file_path)}_table{t_index}_row{r_index}"
                    chunks.append({
                        "chunk_id": chunk_id,
                        "modality": "text",
                        "content": row_text,
                        "metadata": {
                            "file_name": os.path.basename(file_path),
                            "table_index": t_index,
                            "row_index": r_index
                        }
                    })

    except Exception as e:
        print(f"Error processing DOCX file {file_path}: {e}")

    return chunks
