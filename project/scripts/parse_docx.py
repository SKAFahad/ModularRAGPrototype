"""
parse_docx.py

Reads a DOCX file using python-docx, extracting:
  - Paragraph text
  - Table rows
  - Inline images (real files, not placeholders)

Requirements:
    pip install python-docx

We assume you have a subfolder "project/extracted_docx_images" for storing images.
"""

import os
from docx import Document
from typing import List, Dict

def process_docx_file(file_path: str) -> List[Dict]:
    """
    Parses a DOCX file, returning a list of chunk dicts:
      - "text" chunks for each paragraph
      - "table" chunks for each table row
      - "image" chunks for each inline shape (extracted to a real file)

    Args:
        file_path (str): Path to the .docx file
    Returns:
        chunks (List[Dict]): Each dict has 'chunk_id', 'modality', 'content', 'metadata'.
    """
    chunks = []

    # We'll store extracted images in a dedicated folder
    extracted_img_folder = "project/extracted_docx_images"
    if not os.path.exists(extracted_img_folder):
        os.makedirs(extracted_img_folder)

    try:
        doc = Document(file_path)

        # 1) Paragraphs -> text chunks
        for p_i, paragraph in enumerate(doc.paragraphs):
            text_content = paragraph.text.strip()
            if text_content:
                chunk_id = f"{os.path.basename(file_path)}_par_{p_i}"
                chunks.append({
                    "chunk_id": chunk_id,
                    "modality": "text",
                    "content": text_content,
                    "metadata": {
                        "file_name": os.path.basename(file_path),
                        "paragraph_index": p_i
                    }
                })

        # 2) Tables -> table chunks
        for t_i, table in enumerate(doc.tables):
            for r_i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                row_text = ", ".join(cells)
                chunk_id = f"{os.path.basename(file_path)}_table{t_i}_row{r_i}"
                chunks.append({
                    "chunk_id": chunk_id,
                    "modality": "table",
                    "content": row_text,
                    "metadata": {
                        "file_name": os.path.basename(file_path),
                        "table_index": t_i,
                        "row_index": r_i
                    }
                })

        # 3) Inline images
        #    Each shape has an _inline property with a reference to the image relationship.
        for shape_i, inline_shape in enumerate(doc.inline_shapes):
            # We can get the relationship id:
            rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            if not rId:
                continue

            # Access the image part via doc.part.related_parts
            image_part = doc.part.related_parts[rId]
            if not image_part:
                continue

            # Extract extension. Usually docx images might be png, jpeg, etc.
            # If the image_part has no known extension, default to something
            ext = getattr(image_part, 'filename', None)
            if ext:
                ext = os.path.splitext(ext)[1]  # e.g. ".png"
            else:
                ext = ".png"

            # Build an output file path. Example: "MyDoc.docx_img_0.png"
            base_filename = os.path.splitext(os.path.basename(file_path))[0]
            out_filename = f"{base_filename}_img_{shape_i}{ext}"
            out_path = os.path.join(extracted_img_folder, out_filename)

            # image_part.blob is the actual binary data
            with open(out_path, "wb") as f:
                f.write(image_part.blob)

            # Now create a chunk referencing this real image path
            chunk_id = f"{os.path.basename(file_path)}_img_{shape_i}"
            chunks.append({
                "chunk_id": chunk_id,
                "modality": "image",
                "content": out_path,  # the actual image file
                "metadata": {
                    "file_name": os.path.basename(file_path),
                    "image_index": shape_i
                }
            })

    except Exception as e:
        print(f"Error processing DOCX file {file_path}: {e}")

    return chunks
